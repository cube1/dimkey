#!/usr/bin/env python3
"""生成补充缺失敏感类型的 fixture 文件

补充的类型: CreditCard, UkPhone, ZipCode, DriversLicense, Title, IpAddress(增强)
"""
import csv
import os
import sys

import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCENARIO_DIR = os.path.join(os.path.dirname(__file__), "..", "scenarios")

THIN_BORDER = Border(
    left=Side(style="thin"), right=Side(style="thin"),
    top=Side(style="thin"), bottom=Side(style="thin"),
)
HEADER_FILL = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
HEADER_FONT = Font(bold=True, color="FFFFFF", size=11)


# ─────────────────────────────────────────────
# 1. en/uk_customer_records.csv
#    覆盖: UkPhone, CreditCard, UkPostcode, Email
# ─────────────────────────────────────────────
def create_uk_customer_records():
    path = os.path.join(SCENARIO_DIR, "en", "uk_customer_records.csv")
    with open(path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow([
            "Customer ID", "Full Name", "UK Phone", "Email",
            "Credit Card", "Postcode", "Address", "Notes"
        ])
        rows = [
            ["UK-001", "Oliver Thompson", "+44 7911 123456", "oliver.thompson@barclays.co.uk",
             "4539 1488 0343 6467", "SW1A 1AA", "10 Downing Street, London",
             "Premium tier client, renewed contract on 15 March 2024"],
            ["UK-002", "Charlotte Davies", "+44 7700 900123", "charlotte.davies@hsbc.com",
             "5425 2334 1102 8976", "EC2R 8AH", "1 Poultry, City of London",
             "Pending KYC review, passport GB 123456789"],
            ["UK-003", "George Wilson", "07456 789012", "g.wilson@lloyds.co.uk",
             "4716 0912 3456 7890", "M1 1AE", "3 Hardman Street, Manchester",
             "Referred by UK-001, driver's licence WILSO703159GW9IJ"],
            ["UK-004", "Amelia Brown", "+44 20 7946 0958", "amelia.brown@natwest.com",
             "5543 2109 8765 4321", "B1 1BB", "1 Centenary Square, Birmingham",
             "Corporate account, IBAN GB29 NWBK 6016 1331 9268 19"],
            ["UK-005", "Harry Taylor", "07891 234567", "harry.taylor@rbs.co.uk",
             "4929 5678 9012 3456", "EH1 1YZ", "36 St Andrew Square, Edinburgh",
             "Flagged for suspicious transaction on 2024-08-22"],
            ["UK-006", "Isabella Martin", "+44 7534 567890", "isabella.martin@santander.co.uk",
             "5187 6543 2109 8765", "LS1 1UR", "2 Trevelyan Square, Leeds",
             "Student account, ZIP 10001 listed as US correspondence address"],
            ["UK-007", "William Harris", "+44 161 496 0000", "w.harris@virgin.co.uk",
             "4024 0071 2345 6789", "CF10 1EP", "Churchill Way, Cardiff",
             "Joint account with spouse, alt phone 07654 321098"],
            ["UK-008", "Sophia Clark", "07321 098765", "sophia.clark@metro.co.uk",
             "5264 1893 0574 6281", "G1 1DU", "110 Queen Street, Glasgow",
             "Account closed 2024-06-30, credit card cancelled"],
        ]
        for row in rows:
            writer.writerow(row)
    print(f"已生成: {path}")


# ─────────────────────────────────────────────
# 2. en/us_compliance_audit.xlsx
#    覆盖: ZipCode, DriversLicense, CreditCard, SSN, UsPhone
# ─────────────────────────────────────────────
def create_us_compliance_audit():
    path = os.path.join(SCENARIO_DIR, "en", "us_compliance_audit.xlsx")
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Compliance Audit Log"

    headers = [
        "Case ID", "Employee Name", "Title", "SSN", "US Phone",
        "Driver's License", "Credit Card on File", "ZIP Code",
        "Email", "Violation Type", "Status"
    ]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.border = THIN_BORDER
        cell.alignment = Alignment(horizontal="center")

    data = [
        ["AUD-001", "Michael Johnson", "Chief Financial Officer",
         "539-48-2671", "(415) 293-8847", "D123-4567-8901", "4539 1488 0343 6467",
         "94102", "m.johnson@corpx.com", "Unauthorized access", "Open"],
        ["AUD-002", "Jennifer Williams", "VP of Engineering",
         "218-73-9045", "312-547-6183", "W456-7890-1234", "5425 2334 1102 8976",
         "60611", "j.williams@corpx.com", "Data export violation", "Under Review"],
        ["AUD-003", "David Martinez", "Senior Analyst",
         "482-16-3759", "(212) 834-2196", "E789-0123-4567", "4716 0912 3456 7890",
         "10022", "d.martinez@corpx.com", "Policy breach", "Resolved"],
        ["AUD-004", "Sarah Anderson", "Director of Compliance",
         "671-24-8903", "(617) 429-5731", "S012-3456-7890", "5543 2109 8765 4321",
         "02108", "s.anderson@corpx.com", "Insider trading flag", "Escalated"],
        ["AUD-005", "Robert Garcia", "IT Security Manager",
         "345-89-1627", "503-218-4765", "G234-5678-9012", "4929 5678 9012 3456",
         "97205", "r.garcia@corpx.com", "System misconfiguration", "Open"],
        ["AUD-006", "Lisa Chen", "Head of HR",
         "723-51-4089", "(202) 614-3982", "C567-8901-2345", "5187 6543 2109 8765",
         "20006", "l.chen@corpx.com", "Employee data leak", "Under Review"],
        ["AUD-007", "James Thompson", "Software Engineer",
         "156-92-3748", "(310) 872-4591", "T890-1234-5678", "4024 0071 2345 6789",
         "90028", "j.thompson@corpx.com", "Code repository access", "Resolved"],
        ["AUD-008", "Maria Rodriguez", "General Counsel",
         "894-37-5210", "646-903-2178", "R123-4567-8901", "5264 1893 0574 6281",
         "10006", "m.rodriguez@corpx.com", "Conflict of interest", "Open"],
        ["AUD-009", "Kevin White", "Database Administrator",
         "267-84-1935", "(773) 541-8629", "W345-6789-0123", "4485 3210 9876 5432",
         "60606", "k.white@corpx.com", "Privilege escalation", "Escalated"],
        ["AUD-010", "Emily Brown", "Product Manager",
         "418-63-7092", "206-734-5918", "B678-9012-3456", "5391 0567 8901 2345",
         "98101", "e.brown@corpx.com", "Vendor fraud referral", "Open"],
    ]

    for row_idx, row_data in enumerate(data, 2):
        for col_idx, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.border = THIN_BORDER
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    for col in range(1, len(headers) + 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 20

    wb.save(path)
    print(f"已生成: {path}")


# ─────────────────────────────────────────────
# 3. en/international_vendor_contacts.docx
#    覆盖: CreditCard, UkPhone, Passport, IBAN, DriversLicense
# ─────────────────────────────────────────────
def create_international_vendor_contacts():
    path = os.path.join(SCENARIO_DIR, "en", "international_vendor_contacts.docx")
    doc = Document()

    doc.add_heading("International Vendor Contact Directory", level=0)
    doc.add_paragraph(
        "Confidential — For Internal Use Only\n"
        "Last Updated: 2024-11-15\n"
        "Compiled by: Global Procurement Department"
    )

    doc.add_heading("1. European Partners", level=1)
    doc.add_paragraph(
        "Vendor: TechBridge Solutions Ltd (London)\n"
        "Contact: James Fletcher, Senior Account Manager\n"
        "Phone: +44 7911 234567\n"
        "Email: j.fletcher@techbridge.co.uk\n"
        "Passport: 533410987\n"
        "Corporate Credit Card: 4539 7890 1234 5678\n"
        "IBAN: GB82 WEST 1234 5698 7654 32\n"
        "Office: 45 Moorgate, London EC2R 6AR\n"
    )
    doc.add_paragraph(
        "Vendor: München Data GmbH (Germany)\n"
        "Contact: Klaus Weber, Geschäftsführer\n"
        "Phone: +49 89 1234 5678\n"
        "Email: k.weber@muenchen-data.de\n"
        "IBAN: DE89 3704 0044 0532 0130 00\n"
        "Driver's License: B072RRE2I55\n"
        "Office: Maximilianstraße 35, 80539 München\n"
    )
    doc.add_paragraph(
        "Vendor: Paris Analytics SAS (France)\n"
        "Contact: Marie Dupont, Directrice Commerciale\n"
        "Phone: +33 1 42 68 53 00\n"
        "Email: m.dupont@paris-analytics.fr\n"
        "Credit Card: 5425 6789 0123 4567\n"
        "IBAN: FR76 1234 5003 1234 5678 9012 345\n"
        "Passport: 12AB34567\n"
        "Office: 12 Rue de Rivoli, 75001 Paris\n"
    )

    doc.add_heading("2. North American Partners", level=1)
    doc.add_paragraph(
        "Vendor: DataFlow Inc. (San Francisco)\n"
        "Contact: Michael Torres, VP of Sales\n"
        "Phone: (415) 555-0147\n"
        "SSN: 345-89-1627\n"
        "Email: m.torres@dataflow.com\n"
        "Credit Card: 4716 5432 1098 7654\n"
        "Driver's License: D890-1234-5678\n"
        "ZIP: 94105\n"
        "Office: 525 Market Street, Suite 3100, San Francisco, CA\n"
    )
    doc.add_paragraph(
        "Vendor: MapleTech Corp (Toronto)\n"
        "Contact: Sarah O'Brien, Director of Partnerships\n"
        "Phone: +1 416-555-0198\n"
        "Email: s.obrien@mapletech.ca\n"
        "Passport: GA 1234567\n"
        "Credit Card: 5543 8765 4321 0987\n"
        "Office: 100 King Street West, Toronto, ON M5X 1A1\n"
    )

    doc.add_heading("3. Payment & Reimbursement Notes", level=1)
    doc.add_paragraph(
        "All vendor reimbursements for Q4 2024 should be processed through the following:\n\n"
        "- TechBridge: Wire to IBAN GB82 WEST 1234 5698 7654 32, ref: TB-Q4-2024\n"
        "- München Data: IBAN DE89 3704 0044 0532 0130 00, ref: MD-Q4-2024\n"
        "- DataFlow: ACH transfer, routing 121000358, account ending 7654\n"
        "- For credit card charges, verify against card ending 5678 (Fletcher) "
        "and card ending 4567 (Dupont)\n\n"
        "Emergency contact: Procurement hotline +44 20 7946 1234\n"
        "Backup: compliance@globalcorp.com"
    )

    doc.save(path)
    print(f"已生成: {path}")


# ─────────────────────────────────────────────
# 4. txt/IT运维事件报告.txt
#    覆盖: IpAddress(IPv4/IPv6/公网), Email, Phone, Title(NER)
# ─────────────────────────────────────────────
def create_it_ops_report():
    path = os.path.join(SCENARIO_DIR, "txt", "IT运维事件报告.txt")
    content = """XX科技有限公司 IT运维事件报告

报告编号：OPS-2024-0915
日期：2024年9月15日
编制人：陈志远（运维总监），手机：13855556666，邮箱：chenzy@xxtech.com

一、事件概述

2024年9月14日 03:27，监控系统告警：生产环境核心服务异常。运维工程师刘阳（手机：18712345678）值班响应。

二、影响范围

受影响服务器：
- Web前端集群：192.168.1.10, 192.168.1.11, 192.168.1.12
- API网关：10.0.2.100 (内网), 公网IP 203.0.113.50
- 数据库主节点：172.16.0.1, 从节点：172.16.0.2, 172.16.0.3
- 日志服务器(IPv6)：2001:0db8:85a3:0000:0000:8a2e:0370:7334
- CDN边缘节点：198.51.100.23, 198.51.100.24
- 备用DNS：8.8.8.8, 2001:4860:4860::8888

三、处理过程

03:30 运维工程师刘阳登录跳板机(10.0.0.1)排查
03:45 通知技术总监赵明华（手机：13900001234，邮箱：zhaomh@xxtech.com）
04:00 DBA专家孙文博（高级数据库管理员，手机：15800009999）介入数据库排查
04:15 发现数据库主节点172.16.0.1磁盘使用率达98%，触发只读保护
04:30 网络工程师周晓峰（手机：13711112222）检查防火墙规则，确认203.0.113.50端口正常
05:00 扩容完成，服务恢复

四、根因分析

日志分析发现，问题源于每日ETL任务(cron: 03:00)产生的临时文件未清理。
ETL服务器IP：10.0.3.50，IPv6地址：fd00::3:50
任务负责人：前端技术经理杨帆（手机：18600001111，邮箱：yangfan@xxtech.com）

五、改进措施

1. 增加磁盘监控阈值告警（80%预警），由运维总监陈志远负责
2. ETL任务增加临时文件清理步骤，由高级数据库管理员孙文博跟进
3. 灾备方案更新：备用服务器192.168.10.100, 2001:db8::1已就绪

六、审批

编制：陈志远（运维总监）
审核：赵明华（技术总监）
批准：王建国（首席技术官）

报告日期：2024年9月15日
"""
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"已生成: {path}")


# ─────────────────────────────────────────────
# 5. docx/集团高管通讯录.docx
#    覆盖: Title(NER), PersonName, Phone, Email, IdCard
# ─────────────────────────────────────────────
def create_executive_directory():
    path = os.path.join(SCENARIO_DIR, "docx", "集团高管通讯录.docx")
    doc = Document()

    doc.add_heading("华兴实业集团高管通讯录", level=0)
    doc.add_paragraph("机密文件 — 仅限集团内部使用\n更新日期：2024年10月\n编制部门：集团行政中心")

    doc.add_heading("一、集团总部", level=1)
    entries = [
        ("王建国", "集团董事长兼首席执行官", "13900001234", "wangjg@huaxing.com",
         "110101196803151234", "北京市朝阳区建国门外大街甲6号中环世贸中心D座38层"),
        ("李明远", "集团总裁兼首席运营官", "13800005678", "limingyuan@huaxing.com",
         "310101197205201234", "上海市浦东新区陆家嘴环路1000号恒生银行大厦"),
        ("张慧芳", "集团副总裁兼首席财务官", "18600009876", "zhanghuifang@huaxing.com",
         "440301198109121234", "深圳市福田区益田路6009号新世界中心"),
        ("陈大伟", "集团副总裁兼首席技术官", "15900004321", "chendawei@huaxing.com",
         "330106198512081234", "杭州市西湖区文三路478号华星时代广场"),
    ]
    for name, title, phone, email, idcard, addr in entries:
        doc.add_paragraph(
            f"姓名：{name}\n"
            f"职务：{title}\n"
            f"手机：{phone}\n"
            f"邮箱：{email}\n"
            f"身份证号：{idcard}\n"
            f"办公地址：{addr}\n"
        )

    doc.add_heading("二、业务板块负责人", level=1)
    entries2 = [
        ("赵鹏飞", "科技事业部总经理", "13700001111", "zhaopengfei@huaxing.com",
         "510107199008150098"),
        ("孙丽华", "金融事业部总经理", "18500002222", "sunlihua@huaxing.com",
         "320102198706251234"),
        ("周志强", "地产事业部总经理", "15600003333", "zhouzhiqiang@huaxing.com",
         "430102199509121234"),
        ("吴婷婷", "教育事业部总经理", "13600004444", "wutingting@huaxing.com",
         "610102199004189012"),
        ("郑海波", "医疗健康事业部总经理", "17700005555", "zhenghaibol@huaxing.com",
         "350102199202287654"),
    ]
    for name, title, phone, email, idcard in entries2:
        doc.add_paragraph(
            f"姓名：{name}\n"
            f"职务：{title}\n"
            f"手机：{phone}\n"
            f"邮箱：{email}\n"
            f"身份证号：{idcard}\n"
        )

    doc.add_heading("三、集团职能部门负责人", level=1)

    table = doc.add_table(rows=7, cols=5, style='Table Grid')
    table.rows[0].cells[0].text = "部门"
    table.rows[0].cells[1].text = "负责人"
    table.rows[0].cells[2].text = "职务"
    table.rows[0].cells[3].text = "手机"
    table.rows[0].cells[4].text = "邮箱"

    dept_data = [
        ("人力资源部", "林晓燕", "人力资源总监", "13800006666", "linxiaoyan@huaxing.com"),
        ("法务部", "黄国栋", "首席法律顾问", "13900007777", "huangguodong@huaxing.com"),
        ("审计部", "马超", "审计总监", "18700008888", "machao@huaxing.com"),
        ("战略发展部", "何芳", "战略发展总监", "15800009999", "hefang@huaxing.com"),
        ("公共关系部", "许文涛", "公关总监", "13600001010", "xuwentao@huaxing.com"),
        ("信息技术部", "方正", "首席信息官", "18900002020", "fangzheng@huaxing.com"),
    ]
    for i, (dept, name, title, phone, email) in enumerate(dept_data, 1):
        table.rows[i].cells[0].text = dept
        table.rows[i].cells[1].text = name
        table.rows[i].cells[2].text = title
        table.rows[i].cells[3].text = phone
        table.rows[i].cells[4].text = email

    doc.add_heading("四、紧急联络", level=1)
    doc.add_paragraph(
        "集团总机：010-85001234\n"
        "24小时安保热线：010-85001235\n"
        "IT支持热线：400-800-1234\n"
        "集团法务紧急联系：黄国栋 13900007777\n"
        "董事会秘书：钱学军（手机：13500003030，邮箱：qianxuejun@huaxing.com）"
    )

    doc.save(path)
    print(f"已生成: {path}")


if __name__ == "__main__":
    create_uk_customer_records()
    create_us_compliance_audit()
    create_international_vendor_contacts()
    create_it_ops_report()
    create_executive_directory()
    print("\n全部 fixture 文件生成完成！")
