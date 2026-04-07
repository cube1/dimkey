#!/usr/bin/env python3
"""生成英文法律行业场景的测试 fixture 文件"""
import csv
import json
import os
from datetime import datetime

import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = os.path.join(os.path.dirname(__file__), "..", "scenarios")
TODAY = datetime.now().strftime("%Y-%m-%d")


def styled_header(ws, headers, fill_color="1B3A5C"):
    header_fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', wrap_text=True)
        cell.border = thin_border


def write_baseline(fixture_rel_path, expected, output_dir):
    """写 .baseline.json sidecar 文件"""
    fixture_filename = os.path.basename(fixture_rel_path)
    baseline_path = os.path.join(output_dir, fixture_filename + ".baseline.json")
    data = {
        "fixture": fixture_rel_path,
        "generated_by": "dimkey-test-design",
        "generated_at": TODAY,
        "expected": expected,
    }
    with open(baseline_path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    print(f"  基线: {baseline_path}")


# =============================================================================
# 1. law_firm_client_intake.xlsx — 律所客户登记表
# =============================================================================
def create_law_firm_client_intake():
    """US law firm client intake form — 10 rows, covers SSN/UsPhone/Email/ZipCode/PersonName/Address/Title"""
    out_dir = os.path.join(BASE_DIR, "xlsx")
    os.makedirs(out_dir, exist_ok=True)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Client Intake"

    headers = [
        "Case ID", "Client Name", "SSN", "Phone", "Email",
        "Mailing Address", "Zip Code", "Attorney Assigned",
        "Attorney Title", "Intake Date", "Case Type",
    ]
    styled_header(ws, headers)

    data = [
        ["CI-2026-001", "James Anderson", "539-48-2671", "(415) 293-8847", "j.anderson@gmail.com",
         "1420 Market Street, Apt 5B, San Francisco, CA", "94103", "Sarah Mitchell",
         "Senior Partner", "2026-01-15", "Personal Injury"],
        ["CI-2026-002", "Emily Watson", "482-73-1956", "(212) 555-0147", "emily.watson@outlook.com",
         "350 Fifth Avenue, Suite 4200, New York, NY", "10118", "Robert Chen",
         "Managing Partner", "2026-01-22", "Employment Discrimination"],
        ["CI-2026-003", "Michael Torres", "671-34-8820", "(305) 442-9631", "mtorres@yahoo.com",
         "8700 NW 36th Street, Suite 310, Doral, FL", "33166", "Sarah Mitchell",
         "Senior Partner", "2026-02-03", "Medical Malpractice"],
        ["CI-2026-004", "Catherine O'Brien", "328-95-4107", "(617) 738-2201", "cobrien@protonmail.com",
         "200 Clarendon Street, 52nd Floor, Boston, MA", "02116", "David Park",
         "Associate Attorney", "2026-02-10", "Wrongful Termination"],
        ["CI-2026-005", "William Hughes", "754-61-3289", "(312) 997-4455", "w.hughes@icloud.com",
         "233 South Wacker Drive, Suite 8400, Chicago, IL", "60606", "Robert Chen",
         "Managing Partner", "2026-02-18", "Corporate Fraud"],
        ["CI-2026-006", "Jessica Ramirez", "413-87-6542", "(713) 621-3378", "jramirez@hotmail.com",
         "1000 Louisiana Street, Suite 5100, Houston, TX", "77002", "Lisa Yamamoto",
         "Of Counsel", "2026-03-01", "Immigration"],
        ["CI-2026-007", "Daniel Kim", "896-23-5174", "(206) 334-8812", "daniel.kim@live.com",
         "999 Third Avenue, Suite 4600, Seattle, WA", "98104", "David Park",
         "Associate Attorney", "2026-03-05", "Intellectual Property"],
        ["CI-2026-008", "Patricia Sullivan", "265-48-9301", "(202) 555-0193", "p.sullivan@gmail.com",
         "1900 K Street NW, Suite 700, Washington, DC", "20006", "Sarah Mitchell",
         "Senior Partner", "2026-03-12", "Family Law"],
        ["CI-2026-009", "Richard Nakamura", "147-62-8835", "(503) 228-7741", "rnakamura@comcast.net",
         "1120 NW Couch Street, Suite 500, Portland, OR", "97209", "Lisa Yamamoto",
         "Of Counsel", "2026-03-20", "Estate Planning"],
        ["CI-2026-010", "Angela Martinez", "583-19-7064", "(480) 966-2234", "a.martinez@aol.com",
         "2398 East Camelback Road, Suite 1050, Phoenix, AZ", "85016", "Robert Chen",
         "Managing Partner", "2026-03-28", "Real Estate Dispute"],
    ]

    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    for row_idx, row_data in enumerate(data, 2):
        for col_idx, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border

    col_widths = [14, 22, 14, 18, 30, 50, 10, 20, 20, 14, 24]
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w

    filepath = os.path.join(out_dir, "law_firm_client_intake.xlsx")
    wb.save(filepath)
    print(f"已生成: {filepath}")

    # baseline
    expected = [
        # SSN — 10 hard
        {"value": "539-48-2671", "type": "Ssn", "count": 1, "note": "client SSN", "assert": "hard"},
        {"value": "482-73-1956", "type": "Ssn", "count": 1, "note": "client SSN", "assert": "hard"},
        {"value": "671-34-8820", "type": "Ssn", "count": 1, "note": "client SSN", "assert": "hard"},
        {"value": "328-95-4107", "type": "Ssn", "count": 1, "note": "client SSN", "assert": "hard"},
        {"value": "754-61-3289", "type": "Ssn", "count": 1, "note": "client SSN", "assert": "hard"},
        {"value": "413-87-6542", "type": "Ssn", "count": 1, "note": "client SSN", "assert": "hard"},
        {"value": "896-23-5174", "type": "Ssn", "count": 1, "note": "client SSN", "assert": "hard"},
        {"value": "265-48-9301", "type": "Ssn", "count": 1, "note": "client SSN", "assert": "hard"},
        {"value": "147-62-8835", "type": "Ssn", "count": 1, "note": "client SSN", "assert": "hard"},
        {"value": "583-19-7064", "type": "Ssn", "count": 1, "note": "client SSN", "assert": "hard"},
        # UsPhone — 10 hard
        {"value": "(415) 293-8847", "type": "UsPhone", "count": 1, "note": "client phone", "assert": "hard"},
        {"value": "(212) 555-0147", "type": "UsPhone", "count": 1, "note": "client phone", "assert": "hard"},
        {"value": "(305) 442-9631", "type": "UsPhone", "count": 1, "note": "client phone", "assert": "hard"},
        {"value": "(617) 738-2201", "type": "UsPhone", "count": 1, "note": "client phone", "assert": "hard"},
        {"value": "(312) 997-4455", "type": "UsPhone", "count": 1, "note": "client phone", "assert": "hard"},
        {"value": "(713) 621-3378", "type": "UsPhone", "count": 1, "note": "client phone", "assert": "hard"},
        {"value": "(206) 334-8812", "type": "UsPhone", "count": 1, "note": "client phone", "assert": "hard"},
        {"value": "(202) 555-0193", "type": "UsPhone", "count": 1, "note": "client phone", "assert": "hard"},
        {"value": "(503) 228-7741", "type": "UsPhone", "count": 1, "note": "client phone", "assert": "hard"},
        {"value": "(480) 966-2234", "type": "UsPhone", "count": 1, "note": "client phone", "assert": "hard"},
        # Email — 10 hard
        {"value": "j.anderson@gmail.com", "type": "Email", "count": 1, "note": "client email", "assert": "hard"},
        {"value": "emily.watson@outlook.com", "type": "Email", "count": 1, "note": "client email", "assert": "hard"},
        {"value": "mtorres@yahoo.com", "type": "Email", "count": 1, "note": "client email", "assert": "hard"},
        {"value": "cobrien@protonmail.com", "type": "Email", "count": 1, "note": "client email", "assert": "hard"},
        {"value": "w.hughes@icloud.com", "type": "Email", "count": 1, "note": "client email", "assert": "hard"},
        {"value": "jramirez@hotmail.com", "type": "Email", "count": 1, "note": "client email", "assert": "hard"},
        {"value": "daniel.kim@live.com", "type": "Email", "count": 1, "note": "client email", "assert": "hard"},
        {"value": "p.sullivan@gmail.com", "type": "Email", "count": 1, "note": "client email", "assert": "hard"},
        {"value": "rnakamura@comcast.net", "type": "Email", "count": 1, "note": "client email", "assert": "hard"},
        {"value": "a.martinez@aol.com", "type": "Email", "count": 1, "note": "client email", "assert": "hard"},
        # ZipCode — 10 hard
        {"value": "94103", "type": "ZipCode", "count": 1, "note": "mailing zip", "assert": "hard"},
        {"value": "10118", "type": "ZipCode", "count": 1, "note": "mailing zip", "assert": "hard"},
        {"value": "33166", "type": "ZipCode", "count": 1, "note": "mailing zip", "assert": "hard"},
        {"value": "02116", "type": "ZipCode", "count": 1, "note": "mailing zip", "assert": "hard"},
        {"value": "60606", "type": "ZipCode", "count": 1, "note": "mailing zip", "assert": "hard"},
        {"value": "77002", "type": "ZipCode", "count": 1, "note": "mailing zip", "assert": "hard"},
        {"value": "98104", "type": "ZipCode", "count": 1, "note": "mailing zip", "assert": "hard"},
        {"value": "20006", "type": "ZipCode", "count": 1, "note": "mailing zip", "assert": "hard"},
        {"value": "97209", "type": "ZipCode", "count": 1, "note": "mailing zip", "assert": "hard"},
        {"value": "85016", "type": "ZipCode", "count": 1, "note": "mailing zip", "assert": "hard"},
        # PersonName — soft (clients)
        {"value": "James Anderson", "type": "PersonName", "count": 1, "note": "NER client name", "assert": "soft"},
        {"value": "Emily Watson", "type": "PersonName", "count": 1, "note": "NER client name", "assert": "soft"},
        {"value": "Michael Torres", "type": "PersonName", "count": 1, "note": "NER client name", "assert": "soft"},
        {"value": "Catherine O'Brien", "type": "PersonName", "count": 1, "note": "NER client name", "assert": "soft"},
        {"value": "William Hughes", "type": "PersonName", "count": 1, "note": "NER client name", "assert": "soft"},
        {"value": "Jessica Ramirez", "type": "PersonName", "count": 1, "note": "NER client name", "assert": "soft"},
        {"value": "Daniel Kim", "type": "PersonName", "count": 1, "note": "NER client name", "assert": "soft"},
        {"value": "Patricia Sullivan", "type": "PersonName", "count": 1, "note": "NER client name", "assert": "soft"},
        {"value": "Richard Nakamura", "type": "PersonName", "count": 1, "note": "NER client name", "assert": "soft"},
        {"value": "Angela Martinez", "type": "PersonName", "count": 1, "note": "NER client name", "assert": "soft"},
        # PersonName — soft (attorneys)
        {"value": "Sarah Mitchell", "type": "PersonName", "count": 3, "note": "NER attorney name", "assert": "soft"},
        {"value": "Robert Chen", "type": "PersonName", "count": 3, "note": "NER attorney name", "assert": "soft"},
        {"value": "David Park", "type": "PersonName", "count": 2, "note": "NER attorney name", "assert": "soft"},
        {"value": "Lisa Yamamoto", "type": "PersonName", "count": 2, "note": "NER attorney name", "assert": "soft"},
        # Title — soft
        {"value": "Senior Partner", "type": "Title", "count": 3, "note": "NER title", "assert": "soft"},
        {"value": "Managing Partner", "type": "Title", "count": 3, "note": "NER title", "assert": "soft"},
        {"value": "Associate Attorney", "type": "Title", "count": 2, "note": "NER title", "assert": "soft"},
        {"value": "Of Counsel", "type": "Title", "count": 2, "note": "NER title", "assert": "soft"},
        # Address — soft
        {"value": "1420 Market Street, Apt 5B, San Francisco, CA", "type": "Address", "count": 1, "note": "NER address", "assert": "soft"},
        {"value": "350 Fifth Avenue, Suite 4200, New York, NY", "type": "Address", "count": 1, "note": "NER address", "assert": "soft"},
    ]
    write_baseline("scenarios/xlsx/law_firm_client_intake.xlsx", expected, out_dir)


# =============================================================================
# 2. legal_case_management.csv — 案件管理台账
# =============================================================================
def create_legal_case_management():
    """Case management ledger — covers SSN/UsPhone/UkPhone/Email/DriversLicense/Passport/PersonName/OrgName"""
    out_dir = os.path.join(BASE_DIR, "csv")
    os.makedirs(out_dir, exist_ok=True)

    filepath = os.path.join(out_dir, "legal_case_management.csv")
    with open(filepath, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow([
            "Case Number", "Client Name", "Client SSN", "Client Phone", "Client Email",
            "Opposing Party", "Opposing Counsel Phone", "Court",
            "Key Witness", "Witness ID", "Witness Contact", "Status",
        ])
        rows = [
            ["LIT-2026-0401", "Margaret Davies", "412-56-7893", "(202) 445-8821", "m.davies@lawmail.com",
             "Apex Global Holdings LLC", "+44 7700 900123", "US District Court, Eastern District of Virginia",
             "Thomas Whitfield", "D450-3921-8876", "t.whitfield@gmail.com", "Discovery"],
            ["LIT-2026-0402", "Jonathan Reed", "738-21-4560", "(312) 778-3344", "j.reed@outlook.com",
             "Pinnacle Insurance Group", "(617) 902-5567", "Massachusetts Superior Court",
             "Sandra Novak", "E234-5678-9012", "s.novak@yahoo.com", "Pre-Trial"],
            ["LIT-2026-0403", "Olivia Thornton", "295-84-1637", "+44 7456 321987", "o.thornton@barristers.co.uk",
             "British Airways PLC", "+44 20 7946 0958", "High Court of Justice, Queen's Bench",
             "Edward Chambers", "AB123456", "e.chambers@bt.com", "Filed"],
            ["LIT-2026-0404", "Benjamin Cross", "564-39-8712", "(415) 663-2210", "b.cross@protonmail.com",
             "TechVenture Capital Partners", "(650) 555-0198", "San Francisco County Superior Court",
             "Rachel Huang", "F789-0123-4567", "r.huang@stanford.edu", "Mediation"],
            ["LIT-2026-0405", "Charlotte Blake", "821-47-3056", "(713) 445-9932", "c.blake@gmail.com",
             "Gulf Coast Energy Corp", "(281) 334-7721", "Harris County District Court",
             "Marcus Webb", "GA1234567", "m.webb@energy.gov", "Trial"],
            ["LIT-2026-0406", "Alexander Petrov", "673-18-4295", "+44 7911 654321", "a.petrov@solicitors.co.uk",
             "Barclays Bank PLC", "+44 20 3134 5678", "Financial Conduct Authority Tribunal",
             "Natasha Orlova", "GC9876543", "n.orlova@hsbc.com", "Investigation"],
            ["LIT-2026-0407", "Laura Chen", "156-72-9843", "(206) 887-4412", "l.chen@lawfirm.com",
             "Northwest Healthcare Systems", "(503) 221-6654", "King County Superior Court",
             "David Kowalski", "H321-6543-2109", "d.kowalski@uw.edu", "Settlement"],
            ["LIT-2026-0408", "Nathan Brooks", "487-63-2108", "(305) 992-3387", "n.brooks@icloud.com",
             "Caribbean Cruise Lines Inc", "(954) 776-6100", "Miami-Dade Circuit Court",
             "Isabella Vargas", "PA4567890", "i.vargas@cruise.com", "Appeal"],
            ["LIT-2026-0409", "Sophie Hamilton", "342-95-7681", "+44 7800 112233", "s.hamilton@chambers.uk",
             "Lloyd's of London", "+44 20 7327 1000", "Commercial Court, London",
             "George Ashworth", "ASHWO607152AB1CZ", "g.ashworth@lloyds.com", "Hearing"],
            ["LIT-2026-0410", "Christopher Yang", "915-28-6034", "(646) 223-8890", "c.yang@nylaw.com",
             "Manhattan Real Estate Holdings", "(212) 867-5309", "New York Supreme Court",
             "Diane Foster", "K987-6543-2100", "d.foster@realestate.com", "Deposition"],
        ]
        for row in rows:
            writer.writerow(row)
    print(f"已生成: {filepath}")

    expected = [
        # SSN — 10
        {"value": "412-56-7893", "type": "Ssn", "count": 1, "note": "client SSN", "assert": "hard"},
        {"value": "738-21-4560", "type": "Ssn", "count": 1, "note": "client SSN", "assert": "hard"},
        {"value": "295-84-1637", "type": "Ssn", "count": 1, "note": "client SSN", "assert": "hard"},
        {"value": "564-39-8712", "type": "Ssn", "count": 1, "note": "client SSN", "assert": "hard"},
        {"value": "821-47-3056", "type": "Ssn", "count": 1, "note": "client SSN", "assert": "hard"},
        {"value": "673-18-4295", "type": "Ssn", "count": 1, "note": "client SSN", "assert": "hard"},
        {"value": "156-72-9843", "type": "Ssn", "count": 1, "note": "client SSN", "assert": "hard"},
        {"value": "487-63-2108", "type": "Ssn", "count": 1, "note": "client SSN", "assert": "hard"},
        {"value": "342-95-7681", "type": "Ssn", "count": 1, "note": "client SSN", "assert": "hard"},
        {"value": "915-28-6034", "type": "Ssn", "count": 1, "note": "client SSN", "assert": "hard"},
        # UsPhone — 7
        {"value": "(202) 445-8821", "type": "UsPhone", "count": 1, "note": "client phone US", "assert": "hard"},
        {"value": "(312) 778-3344", "type": "UsPhone", "count": 1, "note": "client phone US", "assert": "hard"},
        {"value": "(617) 902-5567", "type": "UsPhone", "count": 1, "note": "opposing counsel phone", "assert": "hard"},
        {"value": "(415) 663-2210", "type": "UsPhone", "count": 1, "note": "client phone US", "assert": "hard"},
        {"value": "(650) 555-0198", "type": "UsPhone", "count": 1, "note": "opposing counsel phone", "assert": "hard"},
        {"value": "(713) 445-9932", "type": "UsPhone", "count": 1, "note": "client phone US", "assert": "hard"},
        {"value": "(281) 334-7721", "type": "UsPhone", "count": 1, "note": "opposing counsel phone", "assert": "hard"},
        {"value": "(206) 887-4412", "type": "UsPhone", "count": 1, "note": "client phone US", "assert": "hard"},
        {"value": "(503) 221-6654", "type": "UsPhone", "count": 1, "note": "opposing counsel phone", "assert": "hard"},
        {"value": "(305) 992-3387", "type": "UsPhone", "count": 1, "note": "client phone US", "assert": "hard"},
        {"value": "(954) 776-6100", "type": "UsPhone", "count": 1, "note": "opposing counsel phone", "assert": "hard"},
        {"value": "(646) 223-8890", "type": "UsPhone", "count": 1, "note": "client phone US", "assert": "hard"},
        {"value": "(212) 867-5309", "type": "UsPhone", "count": 1, "note": "opposing counsel phone", "assert": "hard"},
        # UkPhone — 6
        {"value": "+44 7700 900123", "type": "UkPhone", "count": 1, "note": "opposing counsel UK", "assert": "hard"},
        {"value": "+44 7456 321987", "type": "UkPhone", "count": 1, "note": "client phone UK", "assert": "hard"},
        {"value": "+44 20 7946 0958", "type": "UkPhone", "count": 1, "note": "opposing counsel UK", "assert": "hard"},
        {"value": "+44 7911 654321", "type": "UkPhone", "count": 1, "note": "client phone UK", "assert": "hard"},
        {"value": "+44 20 3134 5678", "type": "UkPhone", "count": 1, "note": "opposing counsel UK", "assert": "hard"},
        {"value": "+44 7800 112233", "type": "UkPhone", "count": 1, "note": "client phone UK", "assert": "hard"},
        {"value": "+44 20 7327 1000", "type": "UkPhone", "count": 1, "note": "opposing counsel UK", "assert": "hard"},
        # Email — 10
        {"value": "m.davies@lawmail.com", "type": "Email", "count": 1, "note": "client email", "assert": "hard"},
        {"value": "j.reed@outlook.com", "type": "Email", "count": 1, "note": "client email", "assert": "hard"},
        {"value": "o.thornton@barristers.co.uk", "type": "Email", "count": 1, "note": "client email", "assert": "hard"},
        {"value": "b.cross@protonmail.com", "type": "Email", "count": 1, "note": "client email", "assert": "hard"},
        {"value": "c.blake@gmail.com", "type": "Email", "count": 1, "note": "client email", "assert": "hard"},
        {"value": "a.petrov@solicitors.co.uk", "type": "Email", "count": 1, "note": "client email", "assert": "hard"},
        {"value": "l.chen@lawfirm.com", "type": "Email", "count": 1, "note": "client email", "assert": "hard"},
        {"value": "n.brooks@icloud.com", "type": "Email", "count": 1, "note": "client email", "assert": "hard"},
        {"value": "s.hamilton@chambers.uk", "type": "Email", "count": 1, "note": "client email", "assert": "hard"},
        {"value": "c.yang@nylaw.com", "type": "Email", "count": 1, "note": "client email", "assert": "hard"},
        {"value": "t.whitfield@gmail.com", "type": "Email", "count": 1, "note": "witness email", "assert": "hard"},
        {"value": "s.novak@yahoo.com", "type": "Email", "count": 1, "note": "witness email", "assert": "hard"},
        {"value": "e.chambers@bt.com", "type": "Email", "count": 1, "note": "witness email", "assert": "hard"},
        {"value": "r.huang@stanford.edu", "type": "Email", "count": 1, "note": "witness email", "assert": "hard"},
        {"value": "m.webb@energy.gov", "type": "Email", "count": 1, "note": "witness email", "assert": "hard"},
        {"value": "n.orlova@hsbc.com", "type": "Email", "count": 1, "note": "witness email", "assert": "hard"},
        {"value": "d.kowalski@uw.edu", "type": "Email", "count": 1, "note": "witness email", "assert": "hard"},
        {"value": "i.vargas@cruise.com", "type": "Email", "count": 1, "note": "witness email", "assert": "hard"},
        {"value": "g.ashworth@lloyds.com", "type": "Email", "count": 1, "note": "witness email", "assert": "hard"},
        {"value": "d.foster@realestate.com", "type": "Email", "count": 1, "note": "witness email", "assert": "hard"},
        # DriversLicense — 5 (US format X123-4567-8901)
        {"value": "D450-3921-8876", "type": "DriversLicense", "count": 1, "note": "US driver's license", "assert": "hard"},
        {"value": "E234-5678-9012", "type": "DriversLicense", "count": 1, "note": "US driver's license", "assert": "hard"},
        {"value": "F789-0123-4567", "type": "DriversLicense", "count": 1, "note": "US driver's license", "assert": "hard"},
        {"value": "H321-6543-2109", "type": "DriversLicense", "count": 1, "note": "US driver's license", "assert": "hard"},
        {"value": "K987-6543-2100", "type": "DriversLicense", "count": 1, "note": "US driver's license", "assert": "hard"},
        # DriversLicense — 1 (UK DVLA format)
        {"value": "ASHWO607152AB1CZ", "type": "DriversLicense", "count": 1, "note": "UK DVLA license", "assert": "hard"},
        # Passport — 4
        {"value": "AB123456", "type": "Passport", "count": 1, "note": "UK passport", "assert": "hard"},
        {"value": "GA1234567", "type": "Passport", "count": 1, "note": "US passport", "assert": "hard"},
        {"value": "GC9876543", "type": "Passport", "count": 1, "note": "UK passport", "assert": "hard"},
        {"value": "PA4567890", "type": "Passport", "count": 1, "note": "US passport", "assert": "hard"},
        # PersonName — soft
        {"value": "Margaret Davies", "type": "PersonName", "count": 1, "note": "NER client name", "assert": "soft"},
        {"value": "Jonathan Reed", "type": "PersonName", "count": 1, "note": "NER client name", "assert": "soft"},
        {"value": "Olivia Thornton", "type": "PersonName", "count": 1, "note": "NER client name", "assert": "soft"},
        {"value": "Benjamin Cross", "type": "PersonName", "count": 1, "note": "NER client name", "assert": "soft"},
        {"value": "Charlotte Blake", "type": "PersonName", "count": 1, "note": "NER client name", "assert": "soft"},
        {"value": "Alexander Petrov", "type": "PersonName", "count": 1, "note": "NER client name", "assert": "soft"},
        {"value": "Laura Chen", "type": "PersonName", "count": 1, "note": "NER client name", "assert": "soft"},
        {"value": "Nathan Brooks", "type": "PersonName", "count": 1, "note": "NER client name", "assert": "soft"},
        {"value": "Sophie Hamilton", "type": "PersonName", "count": 1, "note": "NER client name", "assert": "soft"},
        {"value": "Christopher Yang", "type": "PersonName", "count": 1, "note": "NER client name", "assert": "soft"},
        # OrgName — soft
        {"value": "Apex Global Holdings LLC", "type": "OrgName", "count": 1, "note": "NER opposing party", "assert": "soft"},
        {"value": "Pinnacle Insurance Group", "type": "OrgName", "count": 1, "note": "NER opposing party", "assert": "soft"},
        {"value": "British Airways PLC", "type": "OrgName", "count": 1, "note": "NER opposing party", "assert": "soft"},
        {"value": "Lloyd's of London", "type": "OrgName", "count": 1, "note": "NER opposing party", "assert": "soft"},
    ]
    write_baseline("scenarios/csv/legal_case_management.csv", expected, out_dir)


# =============================================================================
# 3. attorney_engagement_letter.docx — 委托代理协议书
# =============================================================================
def create_attorney_engagement_letter():
    """Engagement letter — covers SSN/UsPhone/Email/CreditCard/ZipCode/PersonName/Address/Title/OrgName"""
    out_dir = os.path.join(BASE_DIR, "docx")
    os.makedirs(out_dir, exist_ok=True)

    doc = Document()

    # Title
    title = doc.add_heading("ATTORNEY-CLIENT ENGAGEMENT AGREEMENT", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Mitchell, Chen & Park LLP\n"
        "2500 Broadway Avenue, Suite 3100\n"
        "San Francisco, CA 94115\n"
        "Tel: (415) 782-3300 | Fax: (415) 782-3301\n"
        "Email: intake@mitchellchenpark.com"
    )

    doc.add_paragraph("Date: March 15, 2026")

    doc.add_paragraph(
        "TO: Mr. James Anderson\n"
        "SSN: 539-48-2671\n"
        "1420 Market Street, Apt 5B\n"
        "San Francisco, CA 94103\n"
        "Phone: (415) 293-8847\n"
        "Email: j.anderson@gmail.com"
    )

    doc.add_heading("1. Scope of Representation", level=1)
    doc.add_paragraph(
        "This engagement letter confirms that Mitchell, Chen & Park LLP (hereinafter \"the Firm\") "
        "has been retained by James Anderson (hereinafter \"the Client\") to provide legal representation "
        "in the matter of Anderson v. Pacific Coast Medical Center (Case No. CV-2026-04587). "
        "The lead attorney on this matter shall be Sarah Mitchell, Senior Partner, with support from "
        "David Park, Associate Attorney."
    )

    doc.add_heading("2. Fee Arrangement", level=1)
    doc.add_paragraph(
        "The Client agrees to pay a retainer of $15,000.00 upon execution of this agreement. "
        "The Firm's hourly rates are as follows:\n\n"
        "- Sarah Mitchell, Senior Partner: $650/hour\n"
        "- David Park, Associate Attorney: $350/hour\n"
        "- Paralegal services: $175/hour\n\n"
        "Payment may be made by credit card: the Client has authorized billing to "
        "Visa card ending in 4539 1488 0343 6467, expiration 09/2028. "
        "Alternatively, payments may be wired to the Firm's IOLTA trust account:\n"
        "IBAN: GB29 NWBK 6016 1331 9268 19\n"
        "Reference: Anderson-CV2026"
    )

    doc.add_heading("3. Client Information", level=1)
    doc.add_paragraph(
        "For conflict-check and identification purposes, the Client provides the following:\n\n"
        "Full Legal Name: James Robert Anderson\n"
        "Date of Birth: April 12, 1985\n"
        "Social Security Number: 539-48-2671\n"
        "Driver's License: D450-3921-8876 (State of California)\n"
        "Employer: TechVenture Capital Partners\n"
        "Work Phone: (650) 555-0198\n"
        "Work Email: janderson@techventure.com"
    )

    doc.add_heading("4. Opposing Party Information", level=1)
    doc.add_paragraph(
        "Defendant: Pacific Coast Medical Center\n"
        "Legal Counsel: Harrison & Associates LLP\n"
        "Contact: Rebecca Harrison, Managing Partner\n"
        "Phone: (415) 331-9920\n"
        "Email: r.harrison@harrisonlaw.com\n"
        "Address: 580 California Street, Suite 2000, San Francisco, CA 94104"
    )

    doc.add_heading("5. Acknowledgment", level=1)
    doc.add_paragraph(
        "By signing below, the Client acknowledges that they have read and understood the terms "
        "of this engagement agreement and consent to the representation described herein."
    )

    doc.add_paragraph("\n\n_________________________          _________________________")
    doc.add_paragraph("James Anderson                      Sarah Mitchell, Senior Partner")
    doc.add_paragraph("Client                              Mitchell, Chen & Park LLP")

    filepath = os.path.join(out_dir, "attorney_engagement_letter.docx")
    doc.save(filepath)
    print(f"已生成: {filepath}")

    expected = [
        # SSN
        {"value": "539-48-2671", "type": "Ssn", "count": 2, "note": "client SSN (appears twice)", "assert": "hard"},
        # UsPhone
        {"value": "(415) 782-3300", "type": "UsPhone", "count": 1, "note": "firm phone", "assert": "hard"},
        {"value": "(415) 782-3301", "type": "UsPhone", "count": 1, "note": "firm fax", "assert": "hard"},
        {"value": "(415) 293-8847", "type": "UsPhone", "count": 1, "note": "client phone", "assert": "hard"},
        {"value": "(650) 555-0198", "type": "UsPhone", "count": 1, "note": "client work phone", "assert": "hard"},
        {"value": "(415) 331-9920", "type": "UsPhone", "count": 1, "note": "opposing counsel phone", "assert": "hard"},
        # Email
        {"value": "intake@mitchellchenpark.com", "type": "Email", "count": 1, "note": "firm email", "assert": "hard"},
        {"value": "j.anderson@gmail.com", "type": "Email", "count": 1, "note": "client email", "assert": "hard"},
        {"value": "janderson@techventure.com", "type": "Email", "count": 1, "note": "client work email", "assert": "hard"},
        {"value": "r.harrison@harrisonlaw.com", "type": "Email", "count": 1, "note": "opposing counsel email", "assert": "hard"},
        # CreditCard
        {"value": "4539 1488 0343 6467", "type": "CreditCard", "count": 1, "note": "Visa card for billing", "assert": "hard"},
        # IBAN
        {"value": "GB29 NWBK 6016 1331 9268 19", "type": "Iban", "count": 1, "note": "IOLTA trust account", "assert": "hard"},
        # DriversLicense
        {"value": "D450-3921-8876", "type": "DriversLicense", "count": 1, "note": "California DL", "assert": "hard"},
        # ZipCode
        {"value": "94115", "type": "ZipCode", "count": 1, "note": "firm zip", "assert": "hard"},
        {"value": "94103", "type": "ZipCode", "count": 1, "note": "client zip", "assert": "hard"},
        {"value": "94104", "type": "ZipCode", "count": 1, "note": "opposing counsel zip", "assert": "hard"},
        # PersonName — soft
        {"value": "James Anderson", "type": "PersonName", "count": 4, "note": "NER client name (multiple mentions)", "assert": "soft"},
        {"value": "Sarah Mitchell", "type": "PersonName", "count": 3, "note": "NER attorney name", "assert": "soft"},
        {"value": "David Park", "type": "PersonName", "count": 1, "note": "NER attorney name", "assert": "soft"},
        {"value": "Rebecca Harrison", "type": "PersonName", "count": 1, "note": "NER opposing counsel", "assert": "soft"},
        # Title — soft
        {"value": "Senior Partner", "type": "Title", "count": 3, "note": "NER title", "assert": "soft"},
        {"value": "Associate Attorney", "type": "Title", "count": 1, "note": "NER title", "assert": "soft"},
        {"value": "Managing Partner", "type": "Title", "count": 1, "note": "NER title", "assert": "soft"},
        # OrgName — soft
        {"value": "Mitchell, Chen & Park LLP", "type": "OrgName", "count": 3, "note": "NER firm name", "assert": "soft"},
        {"value": "Pacific Coast Medical Center", "type": "OrgName", "count": 2, "note": "NER defendant", "assert": "soft"},
        {"value": "Harrison & Associates LLP", "type": "OrgName", "count": 1, "note": "NER opposing firm", "assert": "soft"},
        {"value": "TechVenture Capital Partners", "type": "OrgName", "count": 1, "note": "NER employer", "assert": "soft"},
        # Address — soft
        {"value": "2500 Broadway Avenue, Suite 3100", "type": "Address", "count": 1, "note": "NER firm address", "assert": "soft"},
        {"value": "1420 Market Street, Apt 5B", "type": "Address", "count": 1, "note": "NER client address", "assert": "soft"},
        {"value": "580 California Street, Suite 2000, San Francisco, CA", "type": "Address", "count": 1, "note": "NER opposing address", "assert": "soft"},
    ]
    write_baseline("scenarios/docx/attorney_engagement_letter.docx", expected, out_dir)


# =============================================================================
# 4. litigation_discovery_memo.docx — 诉讼发现备忘录
# =============================================================================
def create_litigation_discovery_memo():
    """Discovery memo — covers SSN/UsPhone/Email/IBAN/Passport/DriversLicense/PersonName/Address"""
    out_dir = os.path.join(BASE_DIR, "docx")
    os.makedirs(out_dir, exist_ok=True)

    doc = Document()

    title = doc.add_heading("PRIVILEGED & CONFIDENTIAL", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("LITIGATION DISCOVERY MEMORANDUM", level=1)

    doc.add_paragraph(
        "Case: Anderson v. Pacific Coast Medical Center\n"
        "Case No: CV-2026-04587\n"
        "Date: March 25, 2026\n"
        "Prepared by: David Park, Associate Attorney\n"
        "Reviewed by: Sarah Mitchell, Senior Partner"
    )

    doc.add_heading("I. Witness List and Contact Information", level=2)

    # Witness table
    table = doc.add_table(rows=7, cols=6, style='Table Grid')
    table.autofit = True
    headers = ["Witness", "Role", "SSN / ID", "Phone", "Email", "Address"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for paragraph in table.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    witnesses = [
        ["Dr. Helen Porter", "Treating Physician", "623-41-8907",
         "(415) 567-2289", "h.porter@pcmedical.org",
         "450 Sutter Street, Suite 800, San Francisco, CA 94108"],
        ["Nurse Linda Chow", "Attending Nurse", "D891-2345-6789",
         "(415) 567-2290", "l.chow@pcmedical.org",
         "1200 Gough Street, Apt 4C, San Francisco, CA 94109"],
        ["Mark Reynolds", "Hospital Administrator", "578-92-3641",
         "(415) 567-2200", "m.reynolds@pcmedical.org",
         "3200 Clay Street, San Francisco, CA 94115"],
        ["Dr. Akira Tanaka", "Expert Witness (Cardiology)", "PA8901234",
         "(310) 445-7823", "a.tanaka@uclahealth.org",
         "10945 Le Conte Avenue, Los Angeles, CA 90095"],
        ["Katherine Wells", "Fact Witness (Spouse)", "487-15-6329",
         "(415) 293-8848", "k.wells@yahoo.com",
         "1420 Market Street, Apt 5B, San Francisco, CA 94103"],
        ["Prof. Richard Eastwood", "Expert Witness (Medical Standards)", "AB987654",
         "+44 20 7123 4567", "r.eastwood@imperial.ac.uk",
         "Exhibition Road, South Kensington, London SW7 2AZ"],
    ]
    for row_idx, w in enumerate(witnesses, 1):
        for col_idx, val in enumerate(w):
            table.rows[row_idx].cells[col_idx].text = val

    doc.add_heading("II. Key Documents Obtained in Discovery", level=2)
    doc.add_paragraph(
        "The following documents have been produced by the Defendant pursuant to our First Request "
        "for Production of Documents:\n\n"
        "1. Medical records for patient James Anderson (SSN: 539-48-2671), spanning January 2025 "
        "through February 2026, totaling 347 pages.\n\n"
        "2. Employment records of Dr. Helen Porter (Employee ID: EMP-2019-0456), including "
        "credentialing documents and malpractice history.\n\n"
        "3. Hospital incident reports filed between December 2025 and February 2026, referencing "
        "the cardiac unit under Dr. Porter's supervision.\n\n"
        "4. Email correspondence between Dr. Porter (h.porter@pcmedical.org) and Hospital Chief "
        "Medical Officer Dr. William Grayson (w.grayson@pcmedical.org) regarding patient complaints."
    )

    doc.add_heading("III. International Expert Consultation", level=2)
    doc.add_paragraph(
        "We have retained Prof. Richard Eastwood of Imperial College London as an independent expert. "
        "His fee arrangement requires wire transfer to the following account:\n\n"
        "Account Holder: Prof. Richard Eastwood\n"
        "IBAN: GB82 WEST 1234 5698 7654 32\n"
        "SWIFT/BIC: WESTGB2L\n"
        "Reference: Anderson-Expert-2026\n\n"
        "Prof. Eastwood's passport number (for travel reimbursement): AB987654\n"
        "UK Driver's License: EASTW753068RE9AZ\n"
        "His assistant, Dr. Elena Markova, can be reached at e.markova@imperial.ac.uk "
        "or +44 7845 123456 for scheduling."
    )

    doc.add_heading("IV. Deposition Schedule", level=2)
    doc.add_paragraph(
        "The following depositions are scheduled:\n\n"
        "1. Dr. Helen Porter — April 10, 2026, 9:00 AM at 580 California Street, Suite 2000\n"
        "   Court reporter: Angela Cruz, (415) 892-3347, a.cruz@sfcourtreporters.com\n\n"
        "2. Mark Reynolds — April 15, 2026, 10:00 AM at same location\n"
        "   Videographer: Jason Liu, (415) 667-4412, j.liu@legalvideo.com\n\n"
        "3. Prof. Richard Eastwood — April 22, 2026, via video conference from London\n"
        "   Remote notary: Emily Saunders, +44 20 7946 0321, e.saunders@uknotary.co.uk"
    )

    filepath = os.path.join(out_dir, "litigation_discovery_memo.docx")
    doc.save(filepath)
    print(f"已生成: {filepath}")

    expected = [
        # SSN
        {"value": "623-41-8907", "type": "Ssn", "count": 1, "note": "witness SSN", "assert": "hard"},
        {"value": "578-92-3641", "type": "Ssn", "count": 1, "note": "witness SSN", "assert": "hard"},
        {"value": "487-15-6329", "type": "Ssn", "count": 1, "note": "witness SSN", "assert": "hard"},
        {"value": "539-48-2671", "type": "Ssn", "count": 1, "note": "patient SSN in docs", "assert": "hard"},
        # UsPhone
        {"value": "(415) 567-2289", "type": "UsPhone", "count": 1, "note": "witness phone", "assert": "hard"},
        {"value": "(415) 567-2290", "type": "UsPhone", "count": 1, "note": "witness phone", "assert": "hard"},
        {"value": "(415) 567-2200", "type": "UsPhone", "count": 1, "note": "witness phone", "assert": "hard"},
        {"value": "(310) 445-7823", "type": "UsPhone", "count": 1, "note": "expert phone", "assert": "hard"},
        {"value": "(415) 293-8848", "type": "UsPhone", "count": 1, "note": "spouse phone", "assert": "hard"},
        {"value": "(415) 892-3347", "type": "UsPhone", "count": 1, "note": "court reporter phone", "assert": "hard"},
        {"value": "(415) 667-4412", "type": "UsPhone", "count": 1, "note": "videographer phone", "assert": "hard"},
        # UkPhone
        {"value": "+44 20 7123 4567", "type": "UkPhone", "count": 1, "note": "UK expert phone", "assert": "hard"},
        {"value": "+44 7845 123456", "type": "UkPhone", "count": 1, "note": "UK assistant phone", "assert": "hard"},
        {"value": "+44 20 7946 0321", "type": "UkPhone", "count": 1, "note": "UK notary phone", "assert": "hard"},
        # Email
        {"value": "h.porter@pcmedical.org", "type": "Email", "count": 2, "note": "witness email (2x)", "assert": "hard"},
        {"value": "l.chow@pcmedical.org", "type": "Email", "count": 1, "note": "witness email", "assert": "hard"},
        {"value": "m.reynolds@pcmedical.org", "type": "Email", "count": 1, "note": "witness email", "assert": "hard"},
        {"value": "a.tanaka@uclahealth.org", "type": "Email", "count": 1, "note": "expert email", "assert": "hard"},
        {"value": "k.wells@yahoo.com", "type": "Email", "count": 1, "note": "spouse email", "assert": "hard"},
        {"value": "r.eastwood@imperial.ac.uk", "type": "Email", "count": 1, "note": "UK expert email", "assert": "hard"},
        {"value": "w.grayson@pcmedical.org", "type": "Email", "count": 1, "note": "CMO email", "assert": "hard"},
        {"value": "e.markova@imperial.ac.uk", "type": "Email", "count": 1, "note": "UK assistant email", "assert": "hard"},
        {"value": "a.cruz@sfcourtreporters.com", "type": "Email", "count": 1, "note": "court reporter email", "assert": "hard"},
        {"value": "j.liu@legalvideo.com", "type": "Email", "count": 1, "note": "videographer email", "assert": "hard"},
        {"value": "e.saunders@uknotary.co.uk", "type": "Email", "count": 1, "note": "UK notary email", "assert": "hard"},
        # IBAN
        {"value": "GB82 WEST 1234 5698 7654 32", "type": "Iban", "count": 1, "note": "expert wire transfer", "assert": "hard"},
        # Passport
        {"value": "PA8901234", "type": "Passport", "count": 1, "note": "expert passport in table", "assert": "hard"},
        {"value": "AB987654", "type": "Passport", "count": 2, "note": "expert passport (table + text)", "assert": "hard"},
        # DriversLicense (US format)
        {"value": "D891-2345-6789", "type": "DriversLicense", "count": 1, "note": "nurse DL as ID", "assert": "hard"},
        # DriversLicense (UK DVLA format)
        {"value": "EASTW753068RE9AZ", "type": "DriversLicense", "count": 1, "note": "UK expert DVLA license", "assert": "hard"},
        # ZipCode
        {"value": "94108", "type": "ZipCode", "count": 1, "note": "witness address zip", "assert": "hard"},
        {"value": "94109", "type": "ZipCode", "count": 1, "note": "witness address zip", "assert": "hard"},
        {"value": "94115", "type": "ZipCode", "count": 1, "note": "witness address zip", "assert": "hard"},
        {"value": "90095", "type": "ZipCode", "count": 1, "note": "expert address zip", "assert": "hard"},
        {"value": "94103", "type": "ZipCode", "count": 1, "note": "spouse address zip", "assert": "hard"},
        # UkPostcode
        {"value": "SW7 2AZ", "type": "UkPostcode", "count": 1, "note": "UK expert address postcode", "assert": "hard"},
        # PersonName — soft
        {"value": "Helen Porter", "type": "PersonName", "count": 4, "note": "NER treating physician", "assert": "soft"},
        {"value": "Linda Chow", "type": "PersonName", "count": 1, "note": "NER nurse", "assert": "soft"},
        {"value": "Mark Reynolds", "type": "PersonName", "count": 2, "note": "NER administrator", "assert": "soft"},
        {"value": "Akira Tanaka", "type": "PersonName", "count": 1, "note": "NER expert", "assert": "soft"},
        {"value": "Katherine Wells", "type": "PersonName", "count": 1, "note": "NER spouse", "assert": "soft"},
        {"value": "Richard Eastwood", "type": "PersonName", "count": 3, "note": "NER UK expert", "assert": "soft"},
        {"value": "James Anderson", "type": "PersonName", "count": 1, "note": "NER patient ref", "assert": "soft"},
        {"value": "David Park", "type": "PersonName", "count": 1, "note": "NER attorney", "assert": "soft"},
        {"value": "Sarah Mitchell", "type": "PersonName", "count": 1, "note": "NER attorney", "assert": "soft"},
        {"value": "William Grayson", "type": "PersonName", "count": 1, "note": "NER CMO", "assert": "soft"},
        {"value": "Elena Markova", "type": "PersonName", "count": 1, "note": "NER UK assistant", "assert": "soft"},
        {"value": "Angela Cruz", "type": "PersonName", "count": 1, "note": "NER court reporter", "assert": "soft"},
        {"value": "Jason Liu", "type": "PersonName", "count": 1, "note": "NER videographer", "assert": "soft"},
        {"value": "Emily Saunders", "type": "PersonName", "count": 1, "note": "NER UK notary", "assert": "soft"},
        # Address — soft
        {"value": "450 Sutter Street, Suite 800, San Francisco, CA", "type": "Address", "count": 1, "note": "NER witness address", "assert": "soft"},
        {"value": "Exhibition Road, South Kensington, London", "type": "Address", "count": 1, "note": "NER UK address", "assert": "soft"},
    ]
    write_baseline("scenarios/docx/litigation_discovery_memo.docx", expected, out_dir)


# =============================================================================
# 5. legal_billing_records.csv — 法律费用账单
# =============================================================================
def create_legal_billing_records():
    """Billing records — covers CreditCard/IBAN/UsPhone/Email/UkPostcode/PersonName/OrgName/Title"""
    out_dir = os.path.join(BASE_DIR, "csv")
    os.makedirs(out_dir, exist_ok=True)

    filepath = os.path.join(out_dir, "legal_billing_records.csv")
    with open(filepath, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow([
            "Invoice No", "Client Name", "Client Org", "Billing Contact",
            "Contact Title", "Phone", "Email", "Payment Method",
            "Card / IBAN", "Billing Address", "Postcode / Zip",
            "Amount (USD)", "Status",
        ])
        rows = [
            ["INV-2026-0101", "James Anderson", "", "James Anderson",
             "CEO", "(415) 293-8847", "j.anderson@gmail.com", "Credit Card",
             "4539 1488 0343 6467", "1420 Market Street, Apt 5B, San Francisco, CA", "94103",
             "15000.00", "Paid"],
            ["INV-2026-0102", "Emily Watson", "Watson & Co Consulting", "Emily Watson",
             "Director", "(212) 555-0147", "emily.watson@outlook.com", "Credit Card",
             "5412 1043 3218 1966", "350 Fifth Avenue, Suite 4200, New York, NY", "10118",
             "8500.00", "Paid"],
            ["INV-2026-0103", "Michael Torres", "Torres Family Trust", "Michael Torres",
             "Trustee", "(305) 442-9631", "mtorres@yahoo.com", "Wire Transfer",
             "GB76 BARC 2026 1799 8843 21", "8700 NW 36th Street, Suite 310, Doral, FL", "33166",
             "22000.00", "Pending"],
            ["INV-2026-0104", "Catherine O'Brien", "O'Brien Ventures Ltd", "Catherine O'Brien",
             "Managing Director", "(617) 738-2201", "cobrien@protonmail.com", "Wire Transfer",
             "GB33 HSBC 4012 7765 4321 09", "200 Clarendon Street, 52nd Floor, Boston, MA", "02116",
             "45000.00", "Paid"],
            ["INV-2026-0105", "Olivia Thornton", "Thornton Chambers", "Olivia Thornton",
             "Barrister", "+44 7456 321987", "o.thornton@barristers.co.uk", "Wire Transfer",
             "GB47 LOYD 3099 1234 5678 90", "15 Old Bailey, London", "EC4M 7EG",
             "18750.00", "Pending"],
            ["INV-2026-0106", "Alexander Petrov", "Petrov & Orlova Solicitors", "Natasha Orlova",
             "Finance Manager", "+44 7911 654321", "n.orlova@hsbc.com", "Wire Transfer",
             "GB91 CITI 1850 0812 3456 78", "1 Churchill Place, Canary Wharf, London", "E14 5HP",
             "31200.00", "Overdue"],
            ["INV-2026-0107", "Laura Chen", "Northwest Healthcare Systems", "Laura Chen",
             "General Counsel", "(206) 887-4412", "l.chen@lawfirm.com", "Credit Card",
             "4916 0013 3890 8380", "999 Third Avenue, Suite 4600, Seattle, WA", "98104",
             "12800.00", "Paid"],
            ["INV-2026-0108", "Nathan Brooks", "", "Nathan Brooks",
             "Attorney", "(305) 992-3387", "n.brooks@icloud.com", "Credit Card",
             "5234 6379 4026 5420", "800 Brickell Avenue, Suite 900, Miami, FL", "33131",
             "9600.00", "Paid"],
            ["INV-2026-0109", "Sophie Hamilton", "Hamilton Chambers", "Sophie Hamilton",
             "Senior Barrister", "+44 7800 112233", "s.hamilton@chambers.uk", "Wire Transfer",
             "GB64 RBOS 8312 4500 9876 54", "4 Paper Buildings, Temple, London", "EC4Y 7EX",
             "27500.00", "Paid"],
            ["INV-2026-0110", "Christopher Yang", "Yang Legal Group PLLC", "Christopher Yang",
             "Founding Partner", "(646) 223-8890", "c.yang@nylaw.com", "Credit Card",
             "4024 3511 6155 9407", "40 Wall Street, 28th Floor, New York, NY", "10005",
             "19400.00", "Pending"],
        ]
        for row in rows:
            writer.writerow(row)
    print(f"已生成: {filepath}")

    expected = [
        # CreditCard — 6 (Visa 4xxx, Mastercard 5xxx)
        {"value": "4539 1488 0343 6467", "type": "CreditCard", "count": 1, "note": "Visa", "assert": "hard"},
        {"value": "5412 1043 3218 1966", "type": "CreditCard", "count": 1, "note": "Mastercard", "assert": "hard"},
        {"value": "4916 0013 3890 8380", "type": "CreditCard", "count": 1, "note": "Visa", "assert": "hard"},
        {"value": "5234 6379 4026 5420", "type": "CreditCard", "count": 1, "note": "Mastercard", "assert": "hard"},
        {"value": "4024 3511 6155 9407", "type": "CreditCard", "count": 1, "note": "Visa", "assert": "hard"},
        # IBAN — 5
        {"value": "GB76 BARC 2026 1799 8843 21", "type": "Iban", "count": 1, "note": "Barclays", "assert": "hard"},
        {"value": "GB33 HSBC 4012 7765 4321 09", "type": "Iban", "count": 1, "note": "HSBC", "assert": "hard"},
        {"value": "GB47 LOYD 3099 1234 5678 90", "type": "Iban", "count": 1, "note": "Lloyds", "assert": "hard"},
        {"value": "GB91 CITI 1850 0812 3456 78", "type": "Iban", "count": 1, "note": "Citi", "assert": "hard"},
        {"value": "GB64 RBOS 8312 4500 9876 54", "type": "Iban", "count": 1, "note": "RBS", "assert": "hard"},
        # UsPhone — 6
        {"value": "(415) 293-8847", "type": "UsPhone", "count": 1, "note": "billing phone", "assert": "hard"},
        {"value": "(212) 555-0147", "type": "UsPhone", "count": 1, "note": "billing phone", "assert": "hard"},
        {"value": "(305) 442-9631", "type": "UsPhone", "count": 1, "note": "billing phone", "assert": "hard"},
        {"value": "(617) 738-2201", "type": "UsPhone", "count": 1, "note": "billing phone", "assert": "hard"},
        {"value": "(206) 887-4412", "type": "UsPhone", "count": 1, "note": "billing phone", "assert": "hard"},
        {"value": "(305) 992-3387", "type": "UsPhone", "count": 1, "note": "billing phone", "assert": "hard"},
        {"value": "(646) 223-8890", "type": "UsPhone", "count": 1, "note": "billing phone", "assert": "hard"},
        # UkPhone — 3
        {"value": "+44 7456 321987", "type": "UkPhone", "count": 1, "note": "UK billing phone", "assert": "hard"},
        {"value": "+44 7911 654321", "type": "UkPhone", "count": 1, "note": "UK billing phone", "assert": "hard"},
        {"value": "+44 7800 112233", "type": "UkPhone", "count": 1, "note": "UK billing phone", "assert": "hard"},
        # Email — 10
        {"value": "j.anderson@gmail.com", "type": "Email", "count": 1, "note": "billing email", "assert": "hard"},
        {"value": "emily.watson@outlook.com", "type": "Email", "count": 1, "note": "billing email", "assert": "hard"},
        {"value": "mtorres@yahoo.com", "type": "Email", "count": 1, "note": "billing email", "assert": "hard"},
        {"value": "cobrien@protonmail.com", "type": "Email", "count": 1, "note": "billing email", "assert": "hard"},
        {"value": "o.thornton@barristers.co.uk", "type": "Email", "count": 1, "note": "billing email", "assert": "hard"},
        {"value": "n.orlova@hsbc.com", "type": "Email", "count": 1, "note": "billing email", "assert": "hard"},
        {"value": "l.chen@lawfirm.com", "type": "Email", "count": 1, "note": "billing email", "assert": "hard"},
        {"value": "n.brooks@icloud.com", "type": "Email", "count": 1, "note": "billing email", "assert": "hard"},
        {"value": "s.hamilton@chambers.uk", "type": "Email", "count": 1, "note": "billing email", "assert": "hard"},
        {"value": "c.yang@nylaw.com", "type": "Email", "count": 1, "note": "billing email", "assert": "hard"},
        # UkPostcode — 3
        {"value": "EC4M 7EG", "type": "UkPostcode", "count": 1, "note": "London postcode", "assert": "hard"},
        {"value": "E14 5HP", "type": "UkPostcode", "count": 1, "note": "Canary Wharf postcode", "assert": "hard"},
        {"value": "EC4Y 7EX", "type": "UkPostcode", "count": 1, "note": "Temple postcode", "assert": "hard"},
        # ZipCode — 7
        {"value": "94103", "type": "ZipCode", "count": 1, "note": "SF zip", "assert": "hard"},
        {"value": "10118", "type": "ZipCode", "count": 1, "note": "NYC zip", "assert": "hard"},
        {"value": "33166", "type": "ZipCode", "count": 1, "note": "Doral zip", "assert": "hard"},
        {"value": "02116", "type": "ZipCode", "count": 1, "note": "Boston zip", "assert": "hard"},
        {"value": "98104", "type": "ZipCode", "count": 1, "note": "Seattle zip", "assert": "hard"},
        {"value": "33131", "type": "ZipCode", "count": 1, "note": "Miami zip", "assert": "hard"},
        {"value": "10005", "type": "ZipCode", "count": 1, "note": "Wall St zip", "assert": "hard"},
        # PersonName — soft
        {"value": "James Anderson", "type": "PersonName", "count": 2, "note": "NER client name", "assert": "soft"},
        {"value": "Emily Watson", "type": "PersonName", "count": 2, "note": "NER client name", "assert": "soft"},
        {"value": "Michael Torres", "type": "PersonName", "count": 2, "note": "NER client name", "assert": "soft"},
        {"value": "Catherine O'Brien", "type": "PersonName", "count": 2, "note": "NER client name", "assert": "soft"},
        {"value": "Olivia Thornton", "type": "PersonName", "count": 2, "note": "NER client name", "assert": "soft"},
        {"value": "Natasha Orlova", "type": "PersonName", "count": 1, "note": "NER contact name", "assert": "soft"},
        {"value": "Laura Chen", "type": "PersonName", "count": 2, "note": "NER client name", "assert": "soft"},
        {"value": "Nathan Brooks", "type": "PersonName", "count": 2, "note": "NER client name", "assert": "soft"},
        {"value": "Sophie Hamilton", "type": "PersonName", "count": 2, "note": "NER client name", "assert": "soft"},
        {"value": "Christopher Yang", "type": "PersonName", "count": 2, "note": "NER client name", "assert": "soft"},
        # OrgName — soft
        {"value": "Watson & Co Consulting", "type": "OrgName", "count": 1, "note": "NER org", "assert": "soft"},
        {"value": "Torres Family Trust", "type": "OrgName", "count": 1, "note": "NER org", "assert": "soft"},
        {"value": "O'Brien Ventures Ltd", "type": "OrgName", "count": 1, "note": "NER org", "assert": "soft"},
        {"value": "Thornton Chambers", "type": "OrgName", "count": 1, "note": "NER org", "assert": "soft"},
        {"value": "Petrov & Orlova Solicitors", "type": "OrgName", "count": 1, "note": "NER org", "assert": "soft"},
        {"value": "Northwest Healthcare Systems", "type": "OrgName", "count": 1, "note": "NER org", "assert": "soft"},
        {"value": "Hamilton Chambers", "type": "OrgName", "count": 1, "note": "NER org", "assert": "soft"},
        {"value": "Yang Legal Group PLLC", "type": "OrgName", "count": 1, "note": "NER org", "assert": "soft"},
        # Title — soft
        {"value": "CEO", "type": "Title", "count": 1, "note": "NER title", "assert": "soft"},
        {"value": "Director", "type": "Title", "count": 1, "note": "NER title", "assert": "soft"},
        {"value": "Trustee", "type": "Title", "count": 1, "note": "NER title", "assert": "soft"},
        {"value": "Managing Director", "type": "Title", "count": 1, "note": "NER title", "assert": "soft"},
        {"value": "Barrister", "type": "Title", "count": 1, "note": "NER title", "assert": "soft"},
        {"value": "Finance Manager", "type": "Title", "count": 1, "note": "NER title", "assert": "soft"},
        {"value": "General Counsel", "type": "Title", "count": 1, "note": "NER title", "assert": "soft"},
        {"value": "Attorney", "type": "Title", "count": 1, "note": "NER title", "assert": "soft"},
        {"value": "Senior Barrister", "type": "Title", "count": 1, "note": "NER title", "assert": "soft"},
        {"value": "Founding Partner", "type": "Title", "count": 1, "note": "NER title", "assert": "soft"},
    ]
    write_baseline("scenarios/csv/legal_billing_records.csv", expected, out_dir)


if __name__ == "__main__":
    create_law_firm_client_intake()
    create_legal_case_management()
    create_attorney_engagement_letter()
    create_litigation_discovery_memo()
    create_legal_billing_records()
    print("\n全部英文法律行业 fixture 生成完毕！")
