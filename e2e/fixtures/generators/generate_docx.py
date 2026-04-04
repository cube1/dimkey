#!/usr/bin/env python3
"""生成用于脱敏工具测试的 Word 文档"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_report_docx():
    """客户调研报告 - 段落中散布各种敏感信息"""
    doc = Document()

    # 标题
    title = doc.add_heading("客户信息调研报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("报告编号: RPT-2024-0315")
    doc.add_paragraph("编制日期: 2024年3月15日")
    doc.add_paragraph("编制人: 市场部 张三")
    doc.add_paragraph("")

    # 第一部分
    doc.add_heading("一、调研背景", level=1)
    doc.add_paragraph(
        "为深入了解核心客户需求，市场部于2024年3月对重点客户进行了走访调研。"
        "本次调研由张三(联系电话: 13800138001，邮箱: zhangsan@qq.com)牵头组织，"
        "共走访客户8家，收集有效反馈32份。"
    )

    # 第二部分
    doc.add_heading("二、走访客户详情", level=1)

    doc.add_heading("2.1 北京星辰科技有限公司", level=2)
    doc.add_paragraph(
        "走访时间: 2024年3月5日\n"
        "对接人: 李四，技术总监\n"
        "联系方式: 手机 15912345678，邮箱 lisi@163.com\n"
        "公司地址: 北京市海淀区中关村大街1号院2号楼\n"
        "李四先生(身份证号: 320106198507121234)对我司产品表示高度认可，"
        "建议增加批量导出功能。付款请转账至对公账户6217001234567890001(工商银行北京分行)。"
    )

    doc.add_heading("2.2 上海云帆信息技术有限公司", level=2)
    doc.add_paragraph(
        "走访时间: 2024年3月8日\n"
        "对接人: 王五，产品经理\n"
        "联系方式: 18676543210，wangwu@gmail.com\n"
        "地址: 上海市浦东新区张江高科技园区碧波路690号\n"
        "王五(身份证440305199212253456)反馈目前使用竞品A，"
        "价格敏感度较高。可通过银行卡6228480012345678901转账。"
    )

    doc.add_heading("2.3 深圳前海智联科技有限公司", level=2)
    doc.add_paragraph(
        "走访时间: 2024年3月10日\n"
        "对接人: 赵六，CEO\n"
        "赵六女士的联系方式是13711112222，邮箱zhaoliu@outlook.com。"
        "公司位于广东省深圳市南山区科技园南区高新南一道008号。"
        "赵六的身份证号为510107199108150098。"
        "她提出希望与我司建立战略合作关系，合同款项可打至6212261234567891234(建设银行深圳支行)。"
    )

    # 第三部分
    doc.add_heading("三、客户反馈汇总", level=1)

    table = doc.add_table(rows=5, cols=5, style='Table Grid')
    headers = ["客户公司", "对接人", "联系电话", "核心诉求", "优先级"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    table_data = [
        ["北京星辰科技有限公司", "李四", "15912345678", "批量导出功能", "高"],
        ["上海云帆信息技术有限公司", "王五", "18676543210", "价格优惠", "中"],
        ["深圳前海智联科技有限公司", "赵六", "13711112222", "战略合作", "高"],
        ["杭州湖畔网络科技有限公司", "陈七", "17088889999", "API对接", "中"],
    ]
    for row_idx, row_data in enumerate(table_data, 1):
        for col_idx, value in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = value

    doc.add_paragraph("")

    # 第四部分
    doc.add_heading("四、结论与建议", level=1)
    doc.add_paragraph(
        "综合以上客户反馈，建议优先开发批量导出功能，并针对上海云帆信息技术有限公司"
        "提供专项优惠方案。后续跟进请联系市场部张三(13800138001)或"
        "发送邮件至zhangsan@qq.com。\n\n"
        "报告人签名: ________________\n"
        "审核人签名: ________________"
    )

    doc.save("/Users/tanzeshun/workpath/git/desensitize-tool/test-data/客户调研报告.docx")
    print("已生成: 客户调研报告.docx")


def create_hr_notice_docx():
    """人事通知 - 包含员工个人敏感信息"""
    doc = Document()

    title = doc.add_heading("人事变动通知", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "各部门负责人：\n\n"
        "经公司管理层研究决定，现对以下人事变动事项进行通知："
    )

    doc.add_heading("一、新员工入职", level=1)
    doc.add_paragraph(
        "1. 吴芳，女，身份证号230102199601080034，"
        "毕业于哈尔滨工业大学计算机科学与技术专业。"
        "手机号14700001111，邮箱wufang@yeah.net。"
        "家庭住址: 黑龙江省哈尔滨市南岗区西大直街92号。"
        "入职日期: 2024年4月1日，分配至研发部，工号EMP009。"
        "工资卡号: 6222081234567890006(中国银行哈尔滨分行)。"
    )
    doc.add_paragraph(
        "2. 郑刚，男，身份证号370102199802145678，"
        "毕业于山东大学软件工程专业。"
        "手机号18300002222，邮箱zhenggang@aliyun.com。"
        "家庭住址: 山东省济南市历下区泉城路180号齐鲁国际大厦。"
        "入职日期: 2024年4月1日，分配至测试部，工号EMP010。"
        "工资卡号: 6217991234567890007(农业银行济南支行)。"
    )

    doc.add_heading("二、岗位调动", level=1)
    doc.add_paragraph(
        "陈七(工号EMP005，手机17088889999)由运维部调至研发部，"
        "担任高级开发工程师一职。调动生效日期: 2024年4月15日。"
        "新工位地址: 杭州市西湖区文三路90号东部软件园A座5楼。"
        "如有工作交接事宜，请联系陈七本人或人力资源部刘八(15233334444，liuba@126.com)。"
    )

    doc.add_heading("三、离职人员", level=1)
    doc.add_paragraph(
        "许婷婷(工号EMP011，身份证号320102198706251234)因个人原因于2024年3月31日正式离职。"
        "最后工资发放至其银行卡6225881234567890002(招商银行南京分行)。"
        "相关工作已移交至孙九(18955556666)负责。"
        "许婷婷原联系邮箱xutingting@tom.com即日起停用。"
    )

    doc.add_paragraph(
        "\n\n特此通知。\n\n"
        "人力资源部\n"
        "北京星辰科技有限公司\n"
        "2024年3月20日"
    )

    doc.save("/Users/tanzeshun/workpath/git/desensitize-tool/test-data/人事变动通知.docx")
    print("已生成: 人事变动通知.docx")


if __name__ == "__main__":
    create_report_docx()
    create_hr_notice_docx()
