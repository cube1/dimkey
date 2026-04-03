#!/usr/bin/env python3
"""生成 E2E 测试用的样本文件，包含各类中文敏感信息"""

import csv
from pathlib import Path

OUTPUT_DIR = Path(__file__).parent

SAMPLE_DATA = [
    ["姓名", "手机号", "身份证号", "邮箱", "地址"],
    ["张三", "13800138000", "110101199001011234", "zhangsan@example.com", "北京市朝阳区建国路88号"],
    ["李四", "13912345678", "320102198512152345", "lisi@test.com", "上海市浦东新区陆家嘴环路1000号"],
    ["王五", "15011112222", "440106197803203456", "wangwu@demo.org", "广州市天河区体育西路191号"],
    ["赵六", "18688889999", "510105199207074567", "zhaoliu@mail.cn", "成都市武侯区人民南路四段1号"],
    ["钱七", "17700001111", "330102198911115678", "qianqi@corp.io", "杭州市西湖区文三路269号"],
]


def generate_csv():
    path = OUTPUT_DIR / "sample.csv"
    with open(path, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerows(SAMPLE_DATA)
    print(f"生成: {path}")


def generate_txt():
    path = OUTPUT_DIR / "sample.txt"
    lines = [
        "员工信息登记表",
        "",
        "张三的手机号是13800138000，身份证号110101199001011234。",
        "联系邮箱：zhangsan@example.com",
        "家庭住址：北京市朝阳区建国路88号",
        "",
        "李四，电话13912345678，身份证320102198512152345。",
        "邮箱：lisi@test.com，住址：上海市浦东新区陆家嘴环路1000号。",
        "",
        "公司名称：阿里巴巴集团控股有限公司",
        "统一社会信用代码：91330100799655058B",
    ]
    path.write_text("\n".join(lines), encoding="utf-8")
    print(f"生成: {path}")


def generate_xlsx():
    try:
        from openpyxl import Workbook
    except ImportError:
        print("跳过 xlsx 生成（需要 openpyxl: pip install openpyxl）")
        return
    wb = Workbook()
    ws = wb.active
    ws.title = "员工信息"
    for row in SAMPLE_DATA:
        ws.append(row)
    path = OUTPUT_DIR / "sample.xlsx"
    wb.save(path)
    print(f"生成: {path}")


def generate_docx():
    try:
        from docx import Document
    except ImportError:
        print("跳过 docx 生成（需要 python-docx: pip install python-docx）")
        return
    doc = Document()
    doc.add_heading("员工信息登记", level=1)
    doc.add_paragraph("张三，手机号13800138000，身份证号110101199001011234。")
    doc.add_paragraph("联系邮箱：zhangsan@example.com")
    doc.add_paragraph("家庭住址：北京市朝阳区建国路88号")
    doc.add_paragraph("")
    doc.add_paragraph("李四，电话13912345678，身份证320102198512152345。")
    doc.add_paragraph("邮箱：lisi@test.com，住址：上海市浦东新区陆家嘴环路1000号。")
    doc.add_paragraph("")
    doc.add_paragraph("公司名称：阿里巴巴集团控股有限公司")
    path = OUTPUT_DIR / "sample.docx"
    doc.save(path)
    print(f"生成: {path}")


if __name__ == "__main__":
    generate_csv()
    generate_txt()
    generate_xlsx()
    generate_docx()
    print("样本文件生成完成")
